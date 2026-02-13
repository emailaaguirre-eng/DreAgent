"""
=============================================================================
HUMMINGBIRD-LEA - Word Document Generation Service
Powered by CoDre-X | B & D Servicing LLC
=============================================================================
Word document generation using python-docx.

Features:
- Professional document layouts
- EIAG-branded documents
- Tables, headers, footers
- Template-based generation
=============================================================================
"""

import logging
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple
import time

from .base import (
    BaseDocumentGenerator,
    DocumentFormat,
    DocumentMetadata,
    DocumentSection,
    DocumentResult,
    ColorScheme,
    FontScheme,
    ProposalData,
    SiteSelectionData,
    IncentiveSummaryData,
)

logger = logging.getLogger(__name__)


# =============================================================================
# Word Document Generator
# =============================================================================

class WordGenerator(BaseDocumentGenerator):
    """
    Word document generator using python-docx.

    Usage:
        generator = WordGenerator()

        # Generate from sections
        result = generator.generate(metadata, sections)

        # Generate EIAG proposal document
        result = generator.generate_proposal(proposal_data)
    """

    @property
    def format(self) -> DocumentFormat:
        return DocumentFormat.DOCX

    def __init__(
        self,
        output_dir: Optional[Path] = None,
        colors: Optional[ColorScheme] = None,
        fonts: Optional[FontScheme] = None,
    ):
        super().__init__(output_dir, colors, fonts)

        # Check for docx availability
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
            self._docx_available = True
        except ImportError:
            self._docx_available = False
            logger.warning("python-docx not installed. Install with: pip install python-docx")

    def _check_docx(self):
        """Check if python-docx is available"""
        if not self._docx_available:
            raise ImportError(
                "python-docx is required for Word document generation. "
                "Install with: pip install python-docx"
            )

    def _hex_to_rgb(self, hex_color: str) -> Tuple[int, int, int]:
        """Convert hex color to RGB tuple"""
        hex_color = hex_color.lstrip("#")
        return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

    def generate(
        self,
        metadata: DocumentMetadata,
        sections: List[DocumentSection],
    ) -> DocumentResult:
        """
        Generate a Word document from sections.

        Args:
            metadata: Document metadata
            sections: List of document sections

        Returns:
            DocumentResult with file path
        """
        self._check_docx()
        start_time = time.time()

        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        try:
            # Create document
            doc = Document()

            # Set document properties
            core_props = doc.core_properties
            core_props.title = metadata.title
            core_props.author = metadata.author
            core_props.subject = metadata.subject
            core_props.keywords = ", ".join(metadata.keywords)
            core_props.category = metadata.category.value

            # Add header
            self._add_header(doc, metadata)

            # Add title page
            self._add_title_page(doc, metadata)

            # Add table of contents placeholder
            self._add_toc_placeholder(doc)

            # Add content sections
            for section in sections:
                self._add_section(doc, section)

            # Add footer
            self._add_footer(doc, metadata)

            # Save file
            filename = self._generate_filename(metadata)
            output_path = self._get_output_path(filename)
            doc.save(str(output_path))

            processing_time = (time.time() - start_time) * 1000

            # Count pages (approximate)
            page_count = max(1, len(sections) // 2 + 1)

            return DocumentResult(
                success=True,
                file_path=output_path,
                format=self.format,
                file_size_bytes=output_path.stat().st_size,
                page_count=page_count,
                processing_time_ms=processing_time,
                metadata=metadata,
            )

        except Exception as e:
            logger.error(f"Word document generation error: {e}")
            return DocumentResult(
                success=False,
                error=str(e),
            )

    def _add_header(self, doc, metadata: DocumentMetadata):
        """Add document header"""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        section = doc.sections[0]
        header = section.header

        # Add company name to header
        para = header.paragraphs[0]
        run = para.add_run(metadata.company)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.colors.text_light))
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    def _add_footer(self, doc, metadata: DocumentMetadata):
        """Add document footer with page numbers"""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        section = doc.sections[0]
        footer = section.footer

        para = footer.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add confidential notice
        run = para.add_run("CONFIDENTIAL | ")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.colors.text_light))

        # Add company
        run = para.add_run(f"{metadata.company}")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.colors.text_light))

    def _add_title_page(self, doc, metadata: DocumentMetadata):
        """Add a title page"""
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Add spacing at top
        for _ in range(5):
            doc.add_paragraph()

        # Add title
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(metadata.title)
        title_run.font.size = Pt(28)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.colors.primary))
        title_run.font.name = self.fonts.heading

        # Add project/client name
        if metadata.project_name or metadata.client_name:
            doc.add_paragraph()
            subtitle_para = doc.add_paragraph()
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_text = metadata.project_name or metadata.client_name
            subtitle_run = subtitle_para.add_run(subtitle_text)
            subtitle_run.font.size = Pt(18)
            subtitle_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.colors.secondary))
            subtitle_run.font.name = self.fonts.body

        # Add date and author
        for _ in range(10):
            doc.add_paragraph()

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_text = metadata.created_at.strftime("%B %d, %Y")
        date_run = date_para.add_run(f"Prepared by: {metadata.author}")
        date_run.font.size = Pt(12)
        date_run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.colors.text_light))

        date_para2 = doc.add_paragraph()
        date_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run2 = date_para2.add_run(date_text)
        date_run2.font.size = Pt(12)
        date_run2.font.color.rgb = RGBColor(*self._hex_to_rgb(self.colors.text_light))

        # Add page break
        doc.add_page_break()

    def _add_toc_placeholder(self, doc):
        """Add a table of contents placeholder"""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        toc_heading = doc.add_heading("Table of Contents", level=1)
        toc_heading.runs[0].font.color.rgb = RGBColor(*self._hex_to_rgb(self.colors.primary))

        # Add placeholder text
        para = doc.add_paragraph()
        run = para.add_run("[Update this table of contents after generating the document]")
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.colors.text_light))

        doc.add_page_break()

    def _add_section(self, doc, section: DocumentSection):
        """Add a section to the document"""
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Add heading based on level
        heading = doc.add_heading(section.title, level=section.level)
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.colors.primary))
            run.font.name = self.fonts.heading

        # Add content
        if section.content:
            para = doc.add_paragraph()
            run = para.add_run(section.content)
            run.font.size = Pt(self.fonts.body_size)
            run.font.name = self.fonts.body
            run.font.color.rgb = RGBColor(*self._hex_to_rgb(self.colors.text))

        # Add bullet items
        if section.items:
            for item in section.items:
                para = doc.add_paragraph(item, style='List Bullet')
                for run in para.runs:
                    run.font.size = Pt(self.fonts.body_size)
                    run.font.name = self.fonts.body

        # Add table
        if section.table_data:
            self._add_table(doc, section.table_data)

        # Add spacing after section
        doc.add_paragraph()

        # Process subsections
        if section.subsections:
            for subsection in section.subsections:
                subsection.level = section.level + 1
                self._add_section(doc, subsection)

    def _add_table(self, doc, table_data: List[List[str]]):
        """Add a table to the document"""
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        if not table_data:
            return

        rows = len(table_data)
        cols = len(table_data[0]) if table_data else 0

        if rows == 0 or cols == 0:
            return

        # Create table
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Fill header row
        header_row = table.rows[0]
        for col_idx, cell_text in enumerate(table_data[0]):
            cell = header_row.cells[col_idx]
            cell.text = str(cell_text)

            # Style header cell
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(255, 255, 255)

            # Set background color (requires XML manipulation)
            shading_elm = parse_xml(
                f'<w:shd {nsdecls("w")} w:fill="{self.colors.primary.lstrip("#")}"/>'
            )
            cell._tc.get_or_add_tcPr().append(shading_elm)

        # Fill data rows
        for row_idx, row_data in enumerate(table_data[1:], start=1):
            row = table.rows[row_idx]
            for col_idx, cell_text in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.text = str(cell_text)

                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = Pt(10)

                # Alternate row colors
                if row_idx % 2 == 0:
                    shading_elm = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="{self.colors.background_alt.lstrip("#")}"/>'
                    )
                    cell._tc.get_or_add_tcPr().append(shading_elm)

    # =========================================================================
    # EIAG Template Methods
    # =========================================================================

    def generate_proposal(self, data: ProposalData) -> DocumentResult:
        """
        Generate an EIAG incentive proposal document.

        Args:
            data: Proposal data

        Returns:
            DocumentResult
        """
        metadata = DocumentMetadata(
            title=f"Economic Incentive Proposal",
            subject=f"Incentive proposal for {data.project_name}",
            client_name=data.client_name,
            project_name=data.project_name,
        )

        sections = [
            DocumentSection(
                title="Executive Summary",
                content=data.executive_summary,
                level=1,
            ),
            DocumentSection(
                title="Project Background",
                content=data.background,
                level=1,
            ),
            DocumentSection(
                title="Proposed Incentives",
                content="The following incentives have been identified for this project:",
                level=1,
                table_data=[
                    ["Incentive Program", "Estimated Value", "Type", "Timeline"],
                    *[
                        [
                            inc.get("name", ""),
                            inc.get("value", ""),
                            inc.get("type", ""),
                            inc.get("timeline", "")
                        ]
                        for inc in data.proposed_incentives
                    ]
                ],
            ),
            DocumentSection(
                title="Financial Summary",
                content="",
                level=1,
                items=[
                    f"Total Estimated Incentive Value: {data.financial_summary.get('total_value', 'TBD')}",
                    f"Required Capital Investment: {data.financial_summary.get('investment', 'TBD')}",
                    f"Net Benefit to Company: {data.financial_summary.get('net_benefit', 'TBD')}",
                    f"Projected ROI: {data.financial_summary.get('roi', 'TBD')}",
                ],
            ),
            DocumentSection(
                title="Implementation Timeline",
                content="The following timeline outlines key milestones:",
                level=1,
                table_data=[
                    ["Phase", "Timeline", "Key Milestone"],
                    *[
                        [
                            item.get("phase", ""),
                            item.get("timeline", ""),
                            item.get("milestone", "")
                        ]
                        for item in data.timeline
                    ]
                ],
            ),
            DocumentSection(
                title="Next Steps",
                content="To proceed with this proposal, we recommend the following actions:",
                level=1,
                items=data.next_steps,
            ),
            DocumentSection(
                title="Contact Information",
                content="",
                level=1,
                items=[
                    f"Contact: {data.contact_info.get('name', '')}",
                    f"Email: {data.contact_info.get('email', '')}",
                    f"Phone: {data.contact_info.get('phone', '')}",
                ],
            ),
        ]

        return self.generate(metadata, sections)

    def generate_site_selection(self, data: SiteSelectionData) -> DocumentResult:
        """
        Generate a site selection report document.

        Args:
            data: Site selection data

        Returns:
            DocumentResult
        """
        metadata = DocumentMetadata(
            title=f"Site Selection Analysis Report",
            subject=f"Site selection for {data.project_name}",
            client_name=data.client_name,
            project_name=data.project_name,
        )

        sections = [
            DocumentSection(
                title="Executive Summary",
                content=data.executive_summary,
                level=1,
            ),
            DocumentSection(
                title="Evaluation Methodology",
                content="Sites were evaluated using the following weighted criteria:",
                level=1,
                table_data=[
                    ["Criterion", "Weight (%)", "Description"],
                    *[
                        [
                            c.get("name", ""),
                            c.get("weight", ""),
                            c.get("description", "")
                        ]
                        for c in data.evaluation_criteria
                    ]
                ],
            ),
            DocumentSection(
                title="Sites Evaluated",
                content="The following sites were evaluated for this project:",
                level=1,
                table_data=[
                    ["Site Name", "Location", "Available Space", "Key Features"],
                    *[
                        [
                            s.get("name", ""),
                            s.get("location", ""),
                            s.get("size", ""),
                            s.get("features", "")
                        ]
                        for s in data.sites_evaluated
                    ]
                ],
            ),
            DocumentSection(
                title="Comparison Matrix",
                content="Score comparison across all evaluation criteria:",
                level=1,
                table_data=data.comparison_matrix,
            ),
            DocumentSection(
                title="Recommendation",
                content=data.recommendation,
                level=1,
            ),
        ]

        # Add detailed analysis for each site
        if data.detailed_analysis:
            detail_section = DocumentSection(
                title="Detailed Site Analysis",
                content="",
                level=1,
                subsections=[],
            )

            for site in data.detailed_analysis:
                site_section = DocumentSection(
                    title=site.get("name", "Site"),
                    content=site.get("analysis", ""),
                    level=2,
                    items=site.get("pros_cons", []),
                )
                detail_section.subsections.append(site_section)

            sections.append(detail_section)

        return self.generate(metadata, sections)

    def generate_incentive_summary(self, data: IncentiveSummaryData) -> DocumentResult:
        """
        Generate a tax credit/incentive summary document.

        Args:
            data: Incentive summary data

        Returns:
            DocumentResult
        """
        metadata = DocumentMetadata(
            title=f"Economic Incentive Summary",
            subject=f"Incentive summary for {data.project_name}",
            client_name=data.client_name,
            project_name=data.project_name,
        )

        sections = [
            DocumentSection(
                title="Overview",
                content=f"Total Estimated Incentive Value: {data.total_incentive_value}",
                level=1,
            ),
            DocumentSection(
                title="Incentive Breakdown",
                content="The following incentives are available for this project:",
                level=1,
                table_data=[
                    ["Program", "Value", "Type", "Duration", "Status"],
                    *[
                        [
                            inc.get("name", ""),
                            inc.get("value", ""),
                            inc.get("type", ""),
                            inc.get("duration", ""),
                            inc.get("status", "")
                        ]
                        for inc in data.incentive_breakdown
                    ]
                ],
            ),
            DocumentSection(
                title="Eligibility Requirements",
                content="The following requirements must be met to qualify:",
                level=1,
                items=data.eligibility_requirements,
            ),
            DocumentSection(
                title="Application Timeline",
                content="",
                level=1,
                table_data=[
                    ["Step", "Timeline", "Action Required"],
                    *[
                        [
                            item.get("step", ""),
                            item.get("timeline", ""),
                            item.get("action", "")
                        ]
                        for item in data.application_timeline
                    ]
                ],
            ),
            DocumentSection(
                title="Compliance Requirements",
                content="Ongoing compliance requirements include:",
                level=1,
                items=data.compliance_requirements,
            ),
        ]

        # Add caveats and assumptions if present
        if data.caveats:
            sections.append(DocumentSection(
                title="Important Caveats",
                content="",
                level=1,
                items=data.caveats,
            ))

        if data.assumptions:
            sections.append(DocumentSection(
                title="Assumptions",
                content="This analysis is based on the following assumptions:",
                level=1,
                items=data.assumptions,
            ))

        sections.append(DocumentSection(
            title="Contact Information",
            content="",
            level=1,
            items=[
                f"Contact: {data.contact_info.get('name', '')}",
                f"Email: {data.contact_info.get('email', '')}",
                f"Phone: {data.contact_info.get('phone', '')}",
            ],
        ))

        return self.generate(metadata, sections)

    def generate_letter(
        self,
        recipient_name: str,
        recipient_title: str,
        recipient_company: str,
        subject: str,
        body_paragraphs: List[str],
        sender_name: str,
        sender_title: str,
    ) -> DocumentResult:
        """
        Generate a formal business letter.

        Args:
            recipient_name: Name of recipient
            recipient_title: Recipient's title
            recipient_company: Recipient's company
            subject: Letter subject
            body_paragraphs: List of paragraphs for letter body
            sender_name: Name of sender
            sender_title: Sender's title

        Returns:
            DocumentResult
        """
        self._check_docx()
        start_time = time.time()

        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        try:
            doc = Document()

            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)

            # Add date
            date_para = doc.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            from datetime import datetime
            date_run = date_para.add_run(datetime.now().strftime("%B %d, %Y"))
            date_run.font.size = Pt(11)

            doc.add_paragraph()

            # Add recipient address
            recipient_para = doc.add_paragraph()
            recipient_para.add_run(f"{recipient_name}\n").font.size = Pt(11)
            recipient_para.add_run(f"{recipient_title}\n").font.size = Pt(11)
            recipient_para.add_run(f"{recipient_company}").font.size = Pt(11)

            doc.add_paragraph()

            # Add subject
            subject_para = doc.add_paragraph()
            subject_run = subject_para.add_run(f"RE: {subject}")
            subject_run.font.size = Pt(11)
            subject_run.bold = True

            doc.add_paragraph()

            # Add salutation
            salutation = doc.add_paragraph()
            salutation.add_run(f"Dear {recipient_name},").font.size = Pt(11)

            doc.add_paragraph()

            # Add body paragraphs
            for para_text in body_paragraphs:
                para = doc.add_paragraph()
                run = para.add_run(para_text)
                run.font.size = Pt(11)
                para.paragraph_format.space_after = Pt(12)

            doc.add_paragraph()

            # Add closing
            closing = doc.add_paragraph()
            closing.add_run("Sincerely,").font.size = Pt(11)

            for _ in range(3):
                doc.add_paragraph()

            # Add signature
            sig_para = doc.add_paragraph()
            sig_para.add_run(f"{sender_name}\n").font.size = Pt(11)
            sig_para.add_run(f"{sender_title}").font.size = Pt(11)

            # Create metadata
            metadata = DocumentMetadata(
                title=f"Letter to {recipient_name}",
                subject=subject,
            )

            # Save file
            filename = self._generate_filename(metadata, suffix="letter")
            output_path = self._get_output_path(filename)
            doc.save(str(output_path))

            processing_time = (time.time() - start_time) * 1000

            return DocumentResult(
                success=True,
                file_path=output_path,
                format=self.format,
                file_size_bytes=output_path.stat().st_size,
                page_count=1,
                processing_time_ms=processing_time,
                metadata=metadata,
            )

        except Exception as e:
            logger.error(f"Letter generation error: {e}")
            return DocumentResult(
                success=False,
                error=str(e),
            )


# =============================================================================
# Factory Function
# =============================================================================

_word_generator: Optional[WordGenerator] = None


def get_word_generator() -> WordGenerator:
    """Get or create the Word generator singleton"""
    global _word_generator
    if _word_generator is None:
        _word_generator = WordGenerator()
    return _word_generator
