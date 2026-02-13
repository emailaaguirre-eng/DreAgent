"""
=============================================================================
HUMMINGBIRD-LEA - Document Loaders
Powered by CoDre-X | B & D Servicing LLC
=============================================================================
Document loaders for various file formats.
Supports PDF, DOCX, TXT, and more.

Each loader extracts text content and metadata from documents.
=============================================================================
"""

import logging
import os
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from enum import Enum

logger = logging.getLogger(__name__)


# =============================================================================
# Enums and Constants
# =============================================================================

class DocumentType(Enum):
    """Supported document types"""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    TXT = "txt"
    MD = "md"
    CSV = "csv"
    JSON = "json"
    UNKNOWN = "unknown"

    @classmethod
    def from_extension(cls, ext: str) -> "DocumentType":
        """Get document type from file extension"""
        ext = ext.lower().lstrip(".")
        mapping = {
            "pdf": cls.PDF,
            "docx": cls.DOCX,
            "doc": cls.DOC,
            "txt": cls.TXT,
            "md": cls.MD,
            "csv": cls.CSV,
            "json": cls.JSON,
        }
        return mapping.get(ext, cls.UNKNOWN)


# =============================================================================
# Data Classes
# =============================================================================

@dataclass
class DocumentMetadata:
    """Metadata about a loaded document"""
    source: str                     # File path or URL
    filename: str
    document_type: DocumentType
    created_at: datetime = field(default_factory=datetime.utcnow)
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    char_count: Optional[int] = None
    author: Optional[str] = None
    title: Optional[str] = None
    extra: Dict[str, Any] = field(default_factory=dict)


@dataclass
class Document:
    """A loaded document with content and metadata"""
    content: str
    metadata: DocumentMetadata

    @property
    def word_count(self) -> int:
        """Get word count"""
        return len(self.content.split())

    @property
    def char_count(self) -> int:
        """Get character count"""
        return len(self.content)


@dataclass
class DocumentChunk:
    """A chunk of a document for embedding"""
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    metadata: DocumentMetadata
    embedding: Optional[List[float]] = None

    @property
    def word_count(self) -> int:
        return len(self.content.split())


# =============================================================================
# Base Loader
# =============================================================================

class BaseLoader(ABC):
    """Abstract base class for document loaders"""

    @abstractmethod
    def load(self, file_path: Union[str, Path]) -> Document:
        """
        Load a document from a file path.

        Args:
            file_path: Path to the document

        Returns:
            Document with content and metadata
        """
        pass

    @abstractmethod
    def supports(self, file_path: Union[str, Path]) -> bool:
        """Check if this loader supports the given file"""
        pass

    def _get_file_info(self, file_path: Path) -> Dict[str, Any]:
        """Get basic file info"""
        stat = file_path.stat()
        return {
            "size_bytes": stat.st_size,
            "modified_at": datetime.fromtimestamp(stat.st_mtime),
            "created_at": datetime.fromtimestamp(stat.st_ctime),
        }


# =============================================================================
# Text Loader
# =============================================================================

class TextLoader(BaseLoader):
    """Loader for plain text files (.txt, .md)"""

    SUPPORTED_EXTENSIONS = {".txt", ".md", ".text", ".markdown"}

    def __init__(self, encoding: str = "utf-8"):
        self.encoding = encoding

    def supports(self, file_path: Union[str, Path]) -> bool:
        path = Path(file_path)
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def load(self, file_path: Union[str, Path]) -> Document:
        """Load a text file"""
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        logger.info(f"Loading text file: {path}")

        # Try different encodings
        encodings = [self.encoding, "utf-8", "latin-1", "cp1252"]
        content = None

        for enc in encodings:
            try:
                content = path.read_text(encoding=enc)
                break
            except UnicodeDecodeError:
                continue

        if content is None:
            raise ValueError(f"Could not decode file with any supported encoding: {path}")

        # Get document type
        doc_type = DocumentType.MD if path.suffix.lower() == ".md" else DocumentType.TXT

        metadata = DocumentMetadata(
            source=str(path.absolute()),
            filename=path.name,
            document_type=doc_type,
            word_count=len(content.split()),
            char_count=len(content),
            extra=self._get_file_info(path),
        )

        return Document(content=content, metadata=metadata)


# =============================================================================
# PDF Loader
# =============================================================================

class PDFLoader(BaseLoader):
    """Loader for PDF files using PyPDF2"""

    SUPPORTED_EXTENSIONS = {".pdf"}

    def supports(self, file_path: Union[str, Path]) -> bool:
        path = Path(file_path)
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def load(self, file_path: Union[str, Path]) -> Document:
        """Load a PDF file"""
        try:
            from PyPDF2 import PdfReader
        except ImportError:
            raise ImportError("PyPDF2 is required for PDF loading. Install with: pip install PyPDF2")

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        logger.info(f"Loading PDF file: {path}")

        reader = PdfReader(str(path))

        # Extract text from all pages
        pages_text = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                pages_text.append(text)

        content = "\n\n".join(pages_text)

        # Extract metadata
        pdf_metadata = reader.metadata or {}
        author = pdf_metadata.get("/Author", None)
        title = pdf_metadata.get("/Title", None)

        metadata = DocumentMetadata(
            source=str(path.absolute()),
            filename=path.name,
            document_type=DocumentType.PDF,
            page_count=len(reader.pages),
            word_count=len(content.split()),
            char_count=len(content),
            author=author,
            title=title,
            extra={
                **self._get_file_info(path),
                "pdf_metadata": {k: str(v) for k, v in pdf_metadata.items()},
            },
        )

        return Document(content=content, metadata=metadata)


# =============================================================================
# Word Document Loader
# =============================================================================

class DocxLoader(BaseLoader):
    """Loader for Word documents (.docx) using python-docx"""

    SUPPORTED_EXTENSIONS = {".docx"}

    def supports(self, file_path: Union[str, Path]) -> bool:
        path = Path(file_path)
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def load(self, file_path: Union[str, Path]) -> Document:
        """Load a Word document"""
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required for DOCX loading. Install with: pip install python-docx")

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        logger.info(f"Loading Word document: {path}")

        doc = DocxDocument(str(path))

        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        content = "\n\n".join(paragraphs)

        # Extract metadata
        core_props = doc.core_properties
        author = core_props.author if hasattr(core_props, "author") else None
        title = core_props.title if hasattr(core_props, "title") else None

        metadata = DocumentMetadata(
            source=str(path.absolute()),
            filename=path.name,
            document_type=DocumentType.DOCX,
            word_count=len(content.split()),
            char_count=len(content),
            author=author,
            title=title,
            extra=self._get_file_info(path),
        )

        return Document(content=content, metadata=metadata)


# =============================================================================
# CSV Loader
# =============================================================================

class CSVLoader(BaseLoader):
    """Loader for CSV files"""

    SUPPORTED_EXTENSIONS = {".csv"}

    def __init__(self, encoding: str = "utf-8"):
        self.encoding = encoding

    def supports(self, file_path: Union[str, Path]) -> bool:
        path = Path(file_path)
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def load(self, file_path: Union[str, Path]) -> Document:
        """Load a CSV file as text"""
        import csv

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        logger.info(f"Loading CSV file: {path}")

        rows = []
        with open(path, "r", encoding=self.encoding, newline="") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(" | ".join(row))

        content = "\n".join(rows)

        metadata = DocumentMetadata(
            source=str(path.absolute()),
            filename=path.name,
            document_type=DocumentType.CSV,
            word_count=len(content.split()),
            char_count=len(content),
            extra={
                **self._get_file_info(path),
                "row_count": len(rows),
            },
        )

        return Document(content=content, metadata=metadata)


# =============================================================================
# JSON Loader
# =============================================================================

class JSONLoader(BaseLoader):
    """Loader for JSON files"""

    SUPPORTED_EXTENSIONS = {".json"}

    def supports(self, file_path: Union[str, Path]) -> bool:
        path = Path(file_path)
        return path.suffix.lower() in self.SUPPORTED_EXTENSIONS

    def load(self, file_path: Union[str, Path]) -> Document:
        """Load a JSON file and convert to text"""
        import json

        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {path}")

        logger.info(f"Loading JSON file: {path}")

        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)

        # Convert to formatted text
        content = self._json_to_text(data)

        metadata = DocumentMetadata(
            source=str(path.absolute()),
            filename=path.name,
            document_type=DocumentType.JSON,
            word_count=len(content.split()),
            char_count=len(content),
            extra=self._get_file_info(path),
        )

        return Document(content=content, metadata=metadata)

    def _json_to_text(self, data: Any, depth: int = 0) -> str:
        """Convert JSON data to readable text"""
        indent = "  " * depth

        if isinstance(data, dict):
            lines = []
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    lines.append(f"{indent}{key}:")
                    lines.append(self._json_to_text(value, depth + 1))
                else:
                    lines.append(f"{indent}{key}: {value}")
            return "\n".join(lines)

        elif isinstance(data, list):
            lines = []
            for i, item in enumerate(data):
                if isinstance(item, (dict, list)):
                    lines.append(f"{indent}- Item {i + 1}:")
                    lines.append(self._json_to_text(item, depth + 1))
                else:
                    lines.append(f"{indent}- {item}")
            return "\n".join(lines)

        else:
            return f"{indent}{data}"


# =============================================================================
# Universal Loader
# =============================================================================

class UniversalLoader:
    """
    Universal document loader that automatically selects the right loader
    based on file extension.
    """

    def __init__(self):
        self.loaders: List[BaseLoader] = [
            TextLoader(),
            PDFLoader(),
            DocxLoader(),
            CSVLoader(),
            JSONLoader(),
        ]

    def load(self, file_path: Union[str, Path]) -> Document:
        """
        Load a document using the appropriate loader.

        Args:
            file_path: Path to the document

        Returns:
            Document with content and metadata

        Raises:
            ValueError: If no loader supports the file type
        """
        path = Path(file_path)

        for loader in self.loaders:
            if loader.supports(path):
                return loader.load(path)

        raise ValueError(
            f"Unsupported file type: {path.suffix}. "
            f"Supported types: .txt, .md, .pdf, .docx, .csv, .json"
        )

    def load_directory(
        self,
        directory: Union[str, Path],
        recursive: bool = True,
    ) -> List[Document]:
        """
        Load all supported documents from a directory.

        Args:
            directory: Path to directory
            recursive: Whether to search subdirectories

        Returns:
            List of loaded documents
        """
        path = Path(directory)

        if not path.is_dir():
            raise NotADirectoryError(f"Not a directory: {path}")

        documents = []
        pattern = "**/*" if recursive else "*"

        for file_path in path.glob(pattern):
            if file_path.is_file():
                try:
                    doc = self.load(file_path)
                    documents.append(doc)
                    logger.info(f"Loaded: {file_path.name}")
                except ValueError:
                    # Skip unsupported files
                    continue
                except Exception as e:
                    logger.warning(f"Failed to load {file_path}: {e}")

        logger.info(f"Loaded {len(documents)} documents from {path}")
        return documents


# =============================================================================
# Factory Function
# =============================================================================

def get_loader() -> UniversalLoader:
    """Get the universal document loader"""
    return UniversalLoader()
